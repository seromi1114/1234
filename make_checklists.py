import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_eduzip_checklist(filename):
    doc = docx.Document()
    
    # Title
    p = doc.add_paragraph()
    run = p.add_run("학습지원 소프트웨어 필수기준 체크리스트(공급자용) - 에듀집 탑재용")
    run.font.size = Pt(16)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('□ 제품/서비스 개요', level=2)
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    row0 = table.rows[0].cells
    row0[0].text = "제품/서비스명"
    row0[1].text = "랜덤 발표자 추출기"
    row0[2].text = "공급자"
    row0[3].text = "-"
    
    row1 = table.rows[1].cells
    row1[0].text = "접속경로"
    row1[1].text = "웹 브라우저"
    row1[1].merge(row1[2]).merge(row1[3])
    
    row2 = table.rows[2].cells
    row2[0].text = "주요 내용 및\n기능·특장점"
    row2[1].text = "◦ 교육 및 수업용 무작위 발표자 추출 웹앱\n◦ 브라우저 로컬 스토리지만 사용하여 개인정보 서버 전송 일절 없음\n◦ 깔끔하고 직관적인 UI 및 윤리가이드 탑재"
    row2[1].merge(row2[2]).merge(row2[3])
    
    doc.add_heading('□ 개인정보보호 기준 충족여부', level=2)
    
    t2 = doc.add_table(rows=10, cols=6)
    t2.style = 'Table Grid'
    
    headers = ["선정기준", "세부 내용", "충족", "미충족", "해당\n없음", "증빙"]
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
        
    items = [
        ("1. 최소처리\n원칙 준수", "1-1. 개인정보가 최소한으로 수집되는가?", "■", "□", "□", "개인정보 처리방침 제2조"),
        ("", "1-2. 개인정보 수집·이용 목적이 기재되어 있는가?", "■", "□", "□", "개인정보 처리방침 제1조"),
        ("", "1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?", "■", "□", "□", "개인정보 처리방침 제3조"),
        ("2. 개인정보\n안전조치 의무", "2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?", "■", "□", "□", "개인정보 처리방침 제6조"),
        ("3. 열람/정정/\n삭제 절차", "3-1. 이용자에게 언제든지 자신의 정보를 열람·정정·삭제·처리정지를 요구할 수 있는 절차가 안내되어 있는가?", "■", "□", "□", "개인정보 처리방침 제5조"),
        ("4. 만14세 미만\n아동 보호", "4-1. 만 14세 미만 아동의 경우 법정대리인 동의 등 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?", "□", "□", "■", "서버 수집 없음\n(해당 없음)"),
        ("5. 보호책임자/\n제3자제공/\n위탁 등", "5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?", "□", "□", "■", "서버 운영 안 함\n(해당 없음)"),
        ("", "5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가? (필요시)", "□", "□", "■", "개인정보 처리방침 제4조\n(해당 없음)"),
        ("", "5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가? (필요시)", "□", "□", "■", "개인정보 처리방침 제4조\n(해당 없음)")
    ]
    
    for i, item in enumerate(items):
        row = t2.rows[i+1].cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]
        row[3].text = item[3]
        row[4].text = item[4]
        row[5].text = item[5]
        
    doc.save(filename)

def create_school_checklist(filename):
    doc = docx.Document()
    
    p = doc.add_paragraph()
    run = p.add_run("서식 2 학습지원 소프트웨어 선정기준 체크리스트 (학교용)")
    run.font.size = Pt(16)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('■ 필수기준', level=2)
    
    t1 = doc.add_table(rows=10, cols=6)
    t1.style = 'Table Grid'
    
    headers = ["선정기준", "세부 내용", "충족", "미충족", "해당\n없음", "증빙자료"]
    for i, h in enumerate(headers):
        t1.rows[0].cells[i].text = h
        
    items1 = [
        ("1. 최소처리\n원칙 준수", "1-1. 개인정보가 최소한으로 수집되는가?", "■", "□", "□", "이름만 로컬 수집"),
        ("", "1-2. 개인정보 수집·이용 목적이 기재되어 있는가?", "■", "□", "□", "방침 기재"),
        ("", "1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?", "■", "□", "□", "방침 기재"),
        ("2. 개인정보\n안전조치 의무", "2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?", "■", "□", "□", "로컬 저장 구조"),
        ("3. 열람/정정/\n삭제 절차", "3-1. 이용자에게 언제든지 자신의 정보를 열람·정정·삭제·처리정지를 요구할 수 있는 절차가 안내되어 있는가?", "■", "□", "□", "앱 내 삭제 기능 제공"),
        ("4. 만14세 미만\n아동 보호", "4-1. 만 14세 미만 아동의 경우 법정대리인 동의 등 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?", "□", "□", "■", "서버 전송 없음"),
        ("5. 보호책임자/\n제3자제공/\n위탁 등", "5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?", "□", "□", "■", "로컬 앱이므로 불필요"),
        ("", "5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가? (필요시)", "□", "□", "■", "제3자 제공 없음"),
        ("", "5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가? (필요시)", "□", "□", "■", "위수탁 없음")
    ]
    
    for i, item in enumerate(items1):
        row = t1.rows[i+1].cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]
        row[3].text = item[3]
        row[4].text = item[4]
        row[5].text = item[5]
        
    doc.add_heading('■ 선택기준', level=2)
    
    t2 = doc.add_table(rows=6, cols=6)
    t2.style = 'Table Grid'
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
        
    items2 = [
        ("1. 교육목표 및\n학생특성 적합성", "1-1. 수업 목표와 학생의 학습 수준에 적합한 내용과 기능을 제공하는가?", "■", "□", "□", "발표 및 참여 유도"),
        ("2. 콘텐츠\n품질 및\n안전성", "2-1. 학습 콘텐츠가 정확하고 신뢰할 수 있으며, 학생 연령에 적합·안전한가?", "■", "□", "□", "로컬 기반 안전 구조"),
        ("3. 사용\n환경 적합성", "3-1. 학교의 기기·네트워크 환경에서 모든 학생이 안정적으로 사용할 수 있는가?", "■", "□", "□", "웹 브라우저 기반"),
        ("4. 접근성\n및 사용성", "4-1. 교사와 학생이 필요한 기능과 자료에 쉽게 접근하고 활용할 수 있는가?", "■", "□", "□", "직관적 UI"),
        ("5. 서비스\n운영 및\n지원체계", "5-1. 이용 안내, 기술 지원, 문의 대응 등 서비스 지원 체계를 갖추고 있는가?", "□", "□", "■", "단일 유틸리티 앱")
    ]
    
    for i, item in enumerate(items2):
        row = t2.rows[i+1].cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]
        row[3].text = item[3]
        row[4].text = item[4]
        row[5].text = item[5]

    doc.save(filename)

if __name__ == '__main__':
    create_eduzip_checklist('에듀집 탑재용 체크리스트(랜덤발표자추출기).docx')
    create_school_checklist('학교용 필수기준 체크리스트(랜덤발표자추출기).docx')
    print("Done")
